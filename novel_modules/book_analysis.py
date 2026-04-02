"""
拆书分析模块
支持上传 TXT/DOCX 文件，解析章节结构，进行多维度 AI 分析
"""

import os
import re
import json
import shutil
from datetime import datetime
from typing import List, Dict, Optional, Tuple

# 尝试导入 docx 库（可选依赖）
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

import backend
from novel_modules.billing import record_tokens


class BookAnalysisManager:
    """拆书分析管理器"""

    def __init__(self, book_name: str):
        self.book_name = book_name
        self.analysis_dir = os.path.join("projects", book_name, "analysis")
        self.uploaded_dir = os.path.join(self.analysis_dir, "uploaded")
        self.parsed_dir = os.path.join(self.analysis_dir, "parsed")

        # 确保目录存在
        os.makedirs(self.uploaded_dir, exist_ok=True)
        os.makedirs(self.parsed_dir, exist_ok=True)

    def get_results_path(self) -> str:
        """获取分析结果文件路径"""
        return os.path.join(self.analysis_dir, "results.json")

    def load_results(self) -> Dict:
        """加载分析结果"""
        path = self.get_results_path()
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"books": [], "last_update": None}

    def save_results(self, results: Dict):
        """保存分析结果"""
        results["last_update"] = datetime.now().isoformat()
        path = self.get_results_path()
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)

    def upload_file(self, file_path: str, original_name: str) -> Dict:
        """
        上传文件到分析目录
        返回: {"success": bool, "message": str, "saved_path": str}
        """
        # 检查文件格式
        ext = os.path.splitext(original_name)[1].lower()
        if ext not in ['.txt', '.docx']:
            return {"success": False, "message": f"不支持的文件格式: {ext}，仅支持 TXT 和 DOCX"}

        # 检查 DOCX 支持情况
        if ext == '.docx' and not HAS_DOCX:
            return {"success": False, "message": "未安装 python-docx 库，无法解析 DOCX 文件。请运行: pip install python-docx"}

        # 生成唯一文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[^\w\u4e00-\u9fff.-]', '_', original_name)
        saved_name = f"{timestamp}_{safe_name}"
        saved_path = os.path.join(self.uploaded_dir, saved_name)

        # 复制文件
        shutil.copy(file_path, saved_path)

        return {
            "success": True,
            "message": "文件上传成功",
            "saved_path": saved_path,
            "original_name": original_name
        }

    def parse_txt(self, file_path: str) -> Dict:
        """
        解析 TXT 文件，自动识别章节
        返回: {"success": bool, "chapters": List[Dict], "total_words": int}
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
            except:
                return {"success": False, "message": "无法识别文件编码，请使用 UTF-8 或 GBK 编码"}

        # 章节匹配模式（支持多种格式）
        patterns = [
            r'^第[一二三四五六七八九十百千万零0-9]+[章节回卷部篇][\s\S]*?(?=第[一二三四五六七八九十百千万零0-9]+[章节回卷部篇]|$)',  # 中文格式
            r'^[第Chapter\s]*[0-9]+[章节回卷部篇]?\s*[\.:：]?\s*[\S]+[\s\S]*?(?=^[第Chapter\s]*[0-9]+[章节回卷部篇]?|$)',  # 数字格式
            r'^Chapter\s+\d+[:\.\s][\s\S]*?(?=^Chapter\s+\d+|$)',  # 英文 Chapter 格式
        ]

        chapters = []
        for pattern in patterns:
            matches = re.findall(pattern, content, re.MULTILINE)
            if len(matches) >= 3:  # 至少找到3章才算有效
                for i, match in enumerate(matches):
                    lines = match.strip().split('\n')
                    title = lines[0].strip() if lines else f"第{i+1}章"
                    text = '\n'.join(lines[1:]) if len(lines) > 1 else match

                    # 提取章节号
                    num_match = re.search(r'[0-9]+|[一二三四五六七八九十百千万零]+', title)
                    chapter_num = num_match.group() if num_match else str(i + 1)

                    # 转换中文数字
                    if chapter_num and all(c in '一二三四五六七八九十百千万零' for c in chapter_num):
                        chapter_num = self._chinese_to_arabic(chapter_num)

                    chapters.append({
                        "num": int(chapter_num) if chapter_num.isdigit() else i + 1,
                        "title": title[:100],  # 限制标题长度
                        "content": text.strip(),
                        "word_count": backend.count_words(text)['total_words']
                    })
                break

        # 如果没有找到章节，将整个文件作为一章
        if not chapters:
            chapters = [{
                "num": 1,
                "title": os.path.basename(file_path),
                "content": content.strip(),
                "word_count": backend.count_words(content)['total_words']
            }]

        total_words = sum(c['word_count'] for c in chapters)

        return {
            "success": True,
            "chapters": chapters,
            "total_words": total_words,
            "chapter_count": len(chapters)
        }

    def parse_docx(self, file_path: str) -> Dict:
        """
        解析 DOCX 文件
        返回: {"success": bool, "chapters": List[Dict], "total_words": int}
        """
        if not HAS_DOCX:
            return {"success": False, "message": "未安装 python-docx 库"}

        try:
            doc = Document(file_path)
        except Exception as e:
            return {"success": False, "message": f"无法读取 DOCX 文件: {str(e)}"}

        chapters = []
        current_chapter = None
        current_content = []
        chapter_num = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测是否为章节标题
            is_chapter_title = (
                re.match(r'^第[一二三四五六七八九十百千万零0-9]+[章节回卷部篇]', text) or
                re.match(r'^Chapter\s*\d+', text, re.IGNORECASE) or
                (para.style.name.startswith('Heading') and len(text) < 50)
            )

            if is_chapter_title:
                # 保存之前的章节
                if current_chapter and current_content:
                    content = '\n'.join(current_content)
                    current_chapter['content'] = content
                    current_chapter['word_count'] = backend.count_words(content)['total_words']
                    chapters.append(current_chapter)

                # 开始新章节
                chapter_num += 1
                current_chapter = {
                    "num": chapter_num,
                    "title": text[:100],
                    "content": "",
                    "word_count": 0
                }
                current_content = []
            else:
                current_content.append(text)

        # 保存最后一个章节
        if current_chapter and current_content:
            content = '\n'.join(current_content)
            current_chapter['content'] = content
            current_chapter['word_count'] = backend.count_words(content)['total_words']
            chapters.append(current_chapter)

        # 如果没有找到章节，将整个文档作为一章
        if not chapters:
            all_text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
            chapters = [{
                "num": 1,
                "title": os.path.basename(file_path),
                "content": all_text,
                "word_count": backend.count_words(all_text)['total_words']
            }]

        total_words = sum(c['word_count'] for c in chapters)

        return {
            "success": True,
            "chapters": chapters,
            "total_words": total_words,
            "chapter_count": len(chapters)
        }

    def _chinese_to_arabic(self, chinese_num: str) -> str:
        """将中文数字转换为阿拉伯数字"""
        mapping = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
            '十': 10, '百': 100, '千': 1000, '万': 10000
        }

        result = 0
        temp = 0

        for char in chinese_num:
            if char in mapping:
                value = mapping[char]
                if value >= 10:
                    if temp == 0:
                        temp = 1
                    result += temp * value
                    temp = 0
                else:
                    temp = value

        result += temp
        return str(result)

    def analyze_book(self, file_path: str, book_title: str,
                     analysis_types: List[str], model: str = None) -> Dict:
        """
        对书籍进行多维度分析

        Args:
            file_path: 上传的文件路径
            book_title: 书籍标题
            analysis_types: 分析类型列表，可选值:
                ['comprehensive', 'structure', 'character', 'language', 'plot']
            model: 使用的模型

        Returns:
            分析结果字典
        """
        # 解析文件
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.txt':
            parse_result = self.parse_txt(file_path)
        elif ext == '.docx':
            parse_result = self.parse_docx(file_path)
        else:
            return {"success": False, "message": "不支持的文件格式"}

        if not parse_result['success']:
            return parse_result

        chapters = parse_result['chapters']

        # 保存解析结果
        parsed_path = os.path.join(self.parsed_dir, f"{book_title}_parsed.json")
        with open(parsed_path, 'w', encoding='utf-8') as f:
            json.dump({
                "book_title": book_title,
                "parse_time": datetime.now().isoformat(),
                "chapter_count": len(chapters),
                "total_words": parse_result['total_words'],
                "chapters": [{"num": c["num"], "title": c["title"],
                             "word_count": c["word_count"]} for c in chapters]
            }, f, ensure_ascii=False, indent=2)

        # 获取分析样本（取前几章和中间章节作为样本）
        sample_chapters = self._get_sample_chapters(chapters, sample_count=5)
        sample_text = '\n\n'.join([f"【{c['title']}】\n{c['content'][:2000]}" for c in sample_chapters])

        # 执行各类分析
        analysis_results = {}

        if 'comprehensive' in analysis_types:
            analysis_results['comprehensive'] = self._analyze_comprehensive(sample_text, model)

        if 'structure' in analysis_types:
            analysis_results['structure'] = self._analyze_structure(chapters, model)

        if 'character' in analysis_types:
            analysis_results['character'] = self._analyze_character(sample_text, model)

        if 'language' in analysis_types:
            analysis_results['language'] = self._analyze_language(sample_text, model)

        if 'plot' in analysis_types:
            analysis_results['plot'] = self._analyze_plot(chapters, model)

        # 保存结果
        results = self.load_results()
        book_result = {
            "id": len(results['books']) + 1,
            "title": book_title,
            "file_path": file_path,
            "parse_info": {
                "chapter_count": len(chapters),
                "total_words": parse_result['total_words']
            },
            "analysis_types": analysis_types,
            "results": analysis_results,
            "created_at": datetime.now().isoformat()
        }
        results['books'].append(book_result)
        self.save_results(results)

        return {
            "success": True,
            "message": "分析完成",
            "result": book_result
        }

    def _get_sample_chapters(self, chapters: List[Dict], sample_count: int = 5) -> List[Dict]:
        """获取样本章节（均匀采样）"""
        if len(chapters) <= sample_count:
            return chapters

        step = len(chapters) // sample_count
        samples = []
        for i in range(sample_count):
            idx = min(i * step, len(chapters) - 1)
            samples.append(chapters[idx])

        return samples

    def _call_ai_for_analysis(self, prompt: str, model: str = None) -> Dict:
        """调用 AI 进行分析"""
        if model is None:
            model = backend.CFG.get('writer_model', 'deepseek-chat')

        try:
            client = backend.get_client()

            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一位专业的小说分析师，擅长从写作技法、结构特点、人物塑造、语言风格等多个维度分析文学作品。请用结构化的方式给出分析结果。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            )

            content = response.choices[0].message.content

            # 记录费用
            record_tokens(
                book_name=self.book_name,
                task_type="book_analysis",
                model=model,
                input_tokens=response.usage.prompt_tokens,
                output_tokens=response.usage.completion_tokens,
                config_pricing=backend.CFG.get('pricing', {})
            )

            return {
                "success": True,
                "content": content,
                "usage": {
                    "input_tokens": response.usage.prompt_tokens,
                    "output_tokens": response.usage.completion_tokens
                }
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def _analyze_comprehensive(self, sample_text: str, model: str = None) -> Dict:
        """综合分析：写作技法、结构特点、创作亮点"""
        prompt = f"""请对以下小说片段进行综合分析，从写作技法、结构特点、创作亮点三个维度进行分析。

小说内容：
{sample_text[:4000]}

请按以下格式输出分析结果：

## 写作技法
（分析作者的叙事手法、描写技巧、修辞运用等）

## 结构特点
（分析小说的结构安排、节奏控制、视角运用等）

## 创作亮点
（总结这部作品的独特之处和可学习的地方）

## 综合评价
（总体评价这部作品的文学价值）

请直接输出分析内容，不要有其他废话。"""

        return self._call_ai_for_analysis(prompt, model)

    def _analyze_structure(self, chapters: List[Dict], model: str = None) -> Dict:
        """结构分析：章节布局、情节推进、叙事节奏"""
        # 构建章节概要
        chapter_summary = "\n".join([
            f"第{c['num']}章: {c['title']} ({c['word_count']}字)"
            for c in chapters[:30]  # 最多30章
        ])

        prompt = f"""请分析以下小说的章节结构：

章节列表：
{chapter_summary}

总章数: {len(chapters)}
总字数: {sum(c['word_count'] for c in chapters)}

请从以下维度分析：

## 章节布局
（分析章节长度分布、分章逻辑等）

## 情节推进
（分析故事发展节奏、高潮设置等）

## 叙事节奏
（分析快慢节奏的交替、阅读体验等）

## 结构建议
（如果这是你的作品，你会如何优化结构）

请直接输出分析内容。"""

        return self._call_ai_for_analysis(prompt, model)

    def _analyze_character(self, sample_text: str, model: str = None) -> Dict:
        """人物分析：角色塑造、性格刻画、关系处理"""
        prompt = f"""请分析以下小说片段中的人物塑造：

小说内容：
{sample_text[:4000]}

请从以下维度分析：

## 主要人物
（识别并分析主要角色的性格特点、形象塑造）

## 人物关系
（分析角色之间的关系和互动）

## 对话艺术
（分析对话的技巧和人物语言特点）

## 人物塑造建议
（总结值得学习的人物刻画技巧）

请直接输出分析内容。"""

        return self._call_ai_for_analysis(prompt, model)

    def _analyze_language(self, sample_text: str, model: str = None) -> Dict:
        """语言分析：文风特色、修辞手法、表达技巧"""
        prompt = f"""请分析以下小说片段的语言风格：

小说内容：
{sample_text[:4000]}

请从以下维度分析：

## 文风特色
（分析整体文风、语言风格特点）

## 修辞手法
（分析运用的修辞手法，如比喻、拟人、排比等）

## 表达技巧
（分析描写手法、叙事语言的特点）

## 语言亮点
（摘录几句精彩的描写或对话，并说明好在哪里）

请直接输出分析内容。"""

        return self._call_ai_for_analysis(prompt, model)

    def _analyze_plot(self, chapters: List[Dict], model: str = None) -> Dict:
        """情节分析：冲突设计、悬念营造、转折处理"""
        # 构建章节内容样本
        sample = '\n\n'.join([
            f"【{c['title']}】\n{c['content'][:1500]}"
            for c in self._get_sample_chapters(chapters, 3)
        ])

        prompt = f"""请分析以下小说片段的情节设计：

小说内容：
{sample[:4500]}

请从以下维度分析：

## 冲突设计
（分析故事的矛盾冲突设置、对立关系等）

## 悬念营造
（分析如何制造悬念、吸引读者继续阅读）

## 转折处理
（分析情节转折的技巧和效果）

## 情节节奏
（分析张弛有度的节奏控制）

## 可借鉴之处
（总结值得学习的情节设计技巧）

请直接输出分析内容。"""

        return self._call_ai_for_analysis(prompt, model)

    def list_analyzed_books(self) -> List[Dict]:
        """获取已分析的书籍列表"""
        results = self.load_results()
        return [
            {
                "id": b["id"],
                "title": b["title"],
                "chapter_count": b["parse_info"]["chapter_count"],
                "total_words": b["parse_info"]["total_words"],
                "analysis_types": b["analysis_types"],
                "created_at": b["created_at"]
            }
            for b in results["books"]
        ]

    def get_analysis_detail(self, book_id: int) -> Optional[Dict]:
        """获取指定书籍的详细分析结果"""
        results = self.load_results()
        for book in results["books"]:
            if book["id"] == book_id:
                return book
        return None

    def delete_analysis(self, book_id: int) -> bool:
        """删除指定书籍的分析结果"""
        results = self.load_results()
        original_count = len(results["books"])
        results["books"] = [b for b in results["books"] if b["id"] != book_id]

        if len(results["books"]) < original_count:
            self.save_results(results)
            return True
        return False


# ==================== UI 辅助函数 ====================

_analysis_managers = {}  # 缓存管理器实例

def get_analysis_manager(book_name: str) -> BookAnalysisManager:
    """获取分析管理器实例（带缓存）"""
    if book_name not in _analysis_managers:
        _analysis_managers[book_name] = BookAnalysisManager(book_name)
    return _analysis_managers[book_name]


# ==================== UI 创建函数 ====================

from nicegui import ui, run
from novel_modules.state import app_state, ui_refs


def create_analysis_ui():
    """创建拆书分析 UI"""

    with ui.column().classes('w-full h-full'):
        # 顶部工具栏
        with ui.row().classes('w-full items-center justify-between p-2 bg-grey-1 shrink-0'):
            ui.label('拆书分析').classes('text-h6')
            ui.button('上传书籍', icon='upload_file', on_click=open_upload_dialog) \
                .props('color=primary size=sm')
            ui.button('刷新列表', icon='refresh', on_click=refresh_analysis_list) \
                .props('flat size=sm')

        # 主内容区 - 左右分栏
        with ui.splitter().classes('w-full flex-grow') as splitter:
            with splitter.before:
                # 左侧：已分析书籍列表
                with ui.card().classes('w-full h-full p-2 rounded-none'):
                    ui.label('已分析书籍').classes('text-subtitle1 font-bold mb-2')
                    with ui.scroll_area().classes('h-[calc(100vh-280px)]'):
                        ui_refs['analysis_list_container'] = ui.column().classes('w-full gap-2')

            with splitter.after:
                # 右侧：分析详情
                with ui.card().classes('w-full h-full p-2 rounded-none'):
                    ui_refs['analysis_container'] = ui.column().classes('w-full h-full')
                    with ui_refs['analysis_container']:
                        ui.label('选择一本书籍查看分析结果').classes('text-grey-6')

        # 初始化列表
        refresh_analysis_list()


def refresh_analysis_list():
    """刷新已分析书籍列表"""
    container = ui_refs.get('analysis_list_container')
    if not container:
        return

    container.clear()

    manager = get_analysis_manager(app_state.current_book_name)
    books = manager.list_analyzed_books()

    with container:
        if not books:
            ui.label('暂无分析记录，请上传书籍').classes('text-grey-6 p-4')

        for book in books:
            with ui.card().classes('w-full p-2 cursor-pointer hover:bg-blue-50') \
                    .on('click', lambda b=book: show_analysis_detail(b['id'])):
                with ui.row().classes('w-full items-center justify-between'):
                    with ui.column().classes('flex-grow'):
                        ui.label(book['title']).classes('font-bold text-sm')
                        with ui.row().classes('text-xs text-grey-6 gap-2'):
                            ui.label(f"{book['chapter_count']}章")
                            ui.label(f"{book['total_words']:,}字")
                    ui.button(icon='delete', on_click=lambda e, b=book: delete_analysis(b['id'], e)) \
                        .props('flat round dense size=sm color=red') \
                        .on('click.stop')


def show_analysis_detail(book_id: int):
    """显示分析详情"""
    container = ui_refs.get('analysis_container')
    if not container:
        return

    container.clear()

    manager = get_analysis_manager(app_state.current_book_name)
    detail = manager.get_analysis_detail(book_id)

    if not detail:
        with container:
            ui.label('未找到分析结果').classes('text-red')
        return

    with container:
        # 标题区
        with ui.row().classes('w-full items-center justify-between mb-4'):
            ui.label(detail['title']).classes('text-h6 font-bold')
            with ui.row().classes('gap-2 text-xs text-grey-6'):
                ui.label(f"{detail['parse_info']['chapter_count']}章")
                ui.label(f"{detail['parse_info']['total_words']:,}字")

        # 分析维度标签
        analysis_types = detail.get('analysis_types', [])
        type_labels = {
            'comprehensive': '综合分析',
            'structure': '结构分析',
            'character': '人物分析',
            'language': '语言分析',
            'plot': '情节分析'
        }

        with ui.tabs().classes('w-full bg-grey-2') as tabs:
            tab_panels_map = {}
            for atype in analysis_types:
                label = type_labels.get(atype, atype)
                tab_panels_map[atype] = ui.tab(label)

        with ui.tab_panels(tabs, value=list(tab_panels_map.values())[0] if tab_panels_map else None) \
                .classes('w-full flex-grow'):
            for atype, tab in tab_panels_map.items():
                with ui.tab_panel(tab).classes('w-full'):
                    result = detail['results'].get(atype, {})
                    if result.get('success'):
                        with ui.scroll_area().classes('h-[calc(100vh-380px)]'):
                            ui.markdown(result.get('content', '无内容')).classes('prose max-w-none')
                    else:
                        ui.label(f"分析失败: {result.get('error', '未知错误')}").classes('text-red')


def delete_analysis(book_id: int, event):
    """删除分析结果"""
    manager = get_analysis_manager(app_state.current_book_name)
    success = manager.delete_analysis(book_id)

    if success:
        ui.notify('删除成功', type='positive')
        refresh_analysis_list()

        # 清空详情显示
        container = ui_refs.get('analysis_container')
        if container:
            container.clear()
            with container:
                ui.label('选择一本书籍查看分析结果').classes('text-grey-6')
    else:
        ui.notify('删除失败', type='negative')


async def open_upload_dialog():
    """打开上传对话框"""
    manager = get_analysis_manager(app_state.current_book_name)

    with ui.dialog() as dialog, ui.card().classes('w-[600px] max-h-[80vh]'):
        ui.label('上传书籍').classes('text-h6 font-bold mb-4')

        # 书籍标题输入
        book_title = ui.input(
            label='书籍标题',
            placeholder='给这本书起个名字'
        ).classes('w-full mb-4')

        # 分析类型选择
        analysis_types = []
        type_options = [
            ('comprehensive', '综合分析', True),
            ('structure', '结构分析', False),
            ('character', '人物分析', False),
            ('language', '语言分析', False),
            ('plot', '情节分析', False)
        ]

        ui.label('分析维度（可多选）').classes('text-sm text-grey-7 mb-2')
        with ui.row().classes('w-full gap-2 mb-4 flex-wrap'):
            for type_id, type_name, default in type_options:
                cb = ui.checkbox(type_name, value=default)
                cb._type_id = type_id
                analysis_types.append(cb)

        # 文件上传
        ui.label('选择文件（支持 TXT、DOCX）').classes('text-sm text-grey-7 mb-2')

        uploaded_file = [None]

        async def handle_upload(e):
            if e.name:
                uploaded_file[0] = e
                ui.notify(f'已选择: {e.name}', type='info')

        ui.upload(
            label='点击或拖拽上传',
            auto_upload=False,
            on_upload=handle_upload,
            max_files=1
        ).classes('w-full').props('accept=.txt,.docx')

        # 操作按钮
        with ui.row().classes('w-full justify-end gap-2 mt-4'):
            ui.button('取消', on_click=dialog.close).props('flat')

            async def start_analysis():
                if not book_title.value:
                    ui.notify('请输入书籍标题', type='warning')
                    return

                if not uploaded_file[0]:
                    ui.notify('请选择要上传的文件', type='warning')
                    return

                selected_types = [cb._type_id for cb in analysis_types if cb.value]
                if not selected_types:
                    ui.notify('请至少选择一个分析维度', type='warning')
                    return

                # 显示进度
                ui.notify('正在分析，请稍候...', type='info', timeout=0)

                try:
                    # 执行上传和分析
                    result = await run.io_bound(
                        manager.analyze_book,
                        uploaded_file[0].path,
                        book_title.value,
                        selected_types
                    )

                    ui.notify.dismiss()

                    if result['success']:
                        ui.notify('分析完成！', type='positive')
                        dialog.close()
                        refresh_analysis_list()
                    else:
                        ui.notify(f"分析失败: {result.get('message', '未知错误')}", type='negative')

                except Exception as e:
                    ui.notify.dismiss()
                    ui.notify(f'发生错误: {str(e)}', type='negative')

            ui.button('开始分析', on_click=start_analysis).props('color=primary')

        dialog.open()